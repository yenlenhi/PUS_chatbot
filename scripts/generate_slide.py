from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_presentation_slide():
    doc = Document()
    
    # Title
    title = doc.add_heading('KIẾN TRÚC HỆ THỐNG QUẢN TRỊ (ADMIN DASHBOARD)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle/Highlight
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Giám sát chặt chẽ - Xử lý thông minh - Truy vết minh bạch')
    run.bold = True
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 102, 204) # Blue
    
    doc.add_paragraph() # Spacer

    # Content
    items = [
        {
            "title": "1. Bảng điều khiển (Dashboard)",
            "func": "Giám sát sức khỏe hệ thống & chỉ số theo thời gian thực.",
            "tech": "Real-time Analytics & Health Monitoring"
        },
        {
            "title": "2. Quản lý tài liệu (Documents)",
            "func": "Số hóa văn bản, tự động đọc hiểu văn bản pháp quy.",
            "tech": "AI OCR (Gemini), Semantic Chunking, Vector Database"
        },
        {
            "title": "3. Lịch sử Chat (Audit)",
            "func": "Lưu trữ & truy vết toàn bộ hội thoại để kiểm tra.",
            "tech": "Telemetry Logging & Audit Trail"
        },
        {
            "title": "4. Quản lý File & Biểu mẫu",
            "func": "Tự động cung cấp đúng biểu mẫu người dùng cần.",
            "tech": "Intent Recognition & Keyword Matching"
        },
        {
            "title": "5. Phản hồi (Feedback)",
            "func": "Thu thập đánh giá để tự động cải thiện chất lượng.",
            "tech": "Quality Improvement Loop"
        }
    ]

    for item in items:
        p = doc.add_paragraph()
        # Item Title
        run_title = p.add_run(item["title"] + ": ")
        run_title.bold = True
        run_title.font.size = Pt(12)
        
        # Function
        run_func = p.add_run("\n   • Chức năng: " + item["func"])
        run_func.font.italic = True
        
        # Tech
        run_tech = p.add_run("\n   • Kỹ thuật: " + item["tech"])
        run_tech.bold = True
        
        # Spacing
        p.paragraph_format.space_after = Pt(12)

    # Footer
    footer = doc.add_paragraph('Kiến trúc hệ thống University Chatbot - PUS')
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(10)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.save('Kien_Truc_He_Thong_Admin.docx')
    print("File saved: Kien_Truc_He_Thong_Admin.docx")

if __name__ == "__main__":
    create_presentation_slide()
